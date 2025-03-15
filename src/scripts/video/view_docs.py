from docx import Document
from pathlib import Path
import sys

def print_doc_content(file_path):
    """Print content of a Word document with structure."""
    print(f"\n{'='*80}\n{file_path.name}\n{'='*80}")
    
    doc = Document(file_path)
    
    # Print paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"\n[P{i}] {para.text}")
    
    # Print tables
    for i, table in enumerate(doc.tables):
        print(f"\n[Table {i}]")
        for row in table.rows:
            print(' | '.join(cell.text for cell in row.cells))

def main():
    docs_dir = Path("data/Videos/Videos_summaries")
    
    # Get first 4 documents sorted by lesson number
    docs = sorted(list(docs_dir.glob("*.docx")))[:4]
    
    for doc_path in docs:
        try:
            print_doc_content(doc_path)
        except Exception as e:
            print(f"Error processing {doc_path}: {e}")

if __name__ == "__main__":
    main() 