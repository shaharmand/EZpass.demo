import os
from pathlib import Path
from docx import Document
import openai
from supabase import create_client
import re
from typing import Dict, List, Tuple
import json
from PIL import Image
import io

# Initialize clients
supabase = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_KEY")
)
openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_lesson_info(filename: str) -> Tuple[int, int, str, str]:
    """Extract lesson number, segment, title and infer subtopic from filename."""
    match = re.search(r'שיעור (\d+)\.(\d+)\s*-?\s*(.*?)\.docx', filename)
    if not match:
        return 0, 0, "", ""
    
    lesson_num = int(match.group(1))
    segment_num = int(match.group(2))
    title = match.group(3).strip()
    
    # Infer subtopic from title or mapping
    # TODO: Replace with actual subtopic mapping
    subtopic = title.split('-')[0].strip() if '-' in title else title
    
    return lesson_num, segment_num, title, subtopic

def get_embedding(text: str) -> List[float]:
    """Get embedding for text using OpenAI's API."""
    response = openai.embeddings.create(
        model="text-embedding-3-small",
        input=text,
        encoding_format="float"
    )
    return response.data[0].embedding

def extract_summary(paragraphs: List[str], max_chars: int = 1000) -> str:
    """Extract a summary from the first few meaningful paragraphs."""
    summary = []
    char_count = 0
    
    for para in paragraphs:
        if not para.strip():
            continue
        
        # Skip very short lines (likely headers)
        if len(para.strip()) < 10:
            continue
            
        summary.append(para.strip())
        char_count += len(para)
        
        if char_count >= max_chars:
            break
    
    return '\n'.join(summary)

def process_paragraph(paragraph) -> Dict:
    """Process a paragraph and return its text and metadata."""
    return {
        'text': paragraph.text,
        'style': paragraph.style.name if paragraph.style else 'Normal',
        'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
    }

def extract_tables(doc: Document) -> List[Dict]:
    """Extract content from tables in the document."""
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            'content': table_data,
            'text': ' '.join([cell for row in table_data for cell in row])
        })
    return tables

def process_document(file_path: Path) -> Dict:
    """Process a Word document with emphasis on structure."""
    doc = Document(file_path)
    lesson_num, segment_num, title, subtopic = extract_lesson_info(file_path.name)
    
    # Extract content
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    
    # Get summary from first few paragraphs
    summary = extract_summary(paragraphs)
    
    # Add table content
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_content.append(row_text)
    
    # Combine all content
    full_content = '\n'.join(paragraphs + table_content)
    
    # Create structured embedding text with emphasis
    embedding_text = f"""
    SUBTOPIC: {subtopic}
    SUBTOPIC: {subtopic}
    TITLE: {title}
    TITLE: {title}
    SUMMARY:
    {summary}
    CONTENT:
    {full_content}
    """
    
    embedding = get_embedding(embedding_text)
    
    return {
        'lesson_number': lesson_num,
        'segment_number': segment_num,
        'title': title,
        'subtopic': subtopic,
        'summary': summary,
        'content': full_content,
        'embedding': embedding
    }

def index_documents(docs_dir: Path):
    """Process all documents and store in Supabase with embeddings."""
    for file_path in docs_dir.glob('*.docx'):
        try:
            print(f"Processing {file_path.name}...")
            content = process_document(file_path)
            
            # Store in Supabase
            result = supabase.table('video_content').insert({
                'lesson_number': content['lesson_number'],
                'segment_number': content['segment_number'],
                'title': content['title'],
                'subtopic': content['subtopic'],
                'summary': content['summary'],
                'content': content['content'],
                'embedding': content['embedding']
            }).execute()
            
            print(f"Indexed {file_path.name}")
            
        except Exception as e:
            print(f"Error processing {file_path.name}: {str(e)}")

if __name__ == "__main__":
    docs_dir = Path("data/Videos/Videos_summaries")
    index_documents(docs_dir) 