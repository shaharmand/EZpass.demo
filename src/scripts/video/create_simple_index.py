import os
from pathlib import Path
from docx import Document
import openai
from supabase import create_client
import re
from typing import Dict, List, Tuple

# Initialize clients
supabase = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_KEY")
)
openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_lesson_info(filename: str) -> Tuple[int, int, str]:
    """Extract lesson number, segment and title from filename."""
    match = re.search(r'שיעור (\d+)\.(\d+)\s*-?\s*(.*?)\.docx', filename)
    if match:
        return int(match.group(1)), int(match.group(2)), match.group(3).strip()
    return 0, 0, ""

def get_embedding(text: str) -> List[float]:
    """Get embedding for text using OpenAI's API."""
    response = openai.embeddings.create(
        model="text-embedding-3-small",
        input=text,
        encoding_format="float"
    )
    return response.data[0].embedding

def process_document(file_path: Path) -> Dict:
    """Process a Word document and prepare it for indexing."""
    doc = Document(file_path)
    lesson_num, segment_num, title = extract_lesson_info(file_path.name)
    
    # Extract all text content
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text.strip())
    
    # Add table content
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                content.append(row_text)
    
    full_text = '\n'.join(content)
    
    # Create embedding for the combined title + content
    embedding_text = f"{title}\n{full_text}"
    embedding = get_embedding(embedding_text)
    
    return {
        'lesson_number': lesson_num,
        'segment_number': segment_num,
        'title': title,
        'content': full_text,
        'embedding': embedding
    }

def index_documents(docs_dir: Path):
    """Process all documents and store in Supabase with embeddings."""
    for file_path in docs_dir.glob('*.docx'):
        try:
            print(f"Processing {file_path.name}...")
            content = process_document(file_path)
            
            # Store in Supabase
            result = supabase.table('video_content').insert({
                'lesson_number': content['lesson_number'],
                'segment_number': content['segment_number'],
                'title': content['title'],
                'content': content['content'],
                'embedding': content['embedding']
            }).execute()
            
            print(f"Indexed {file_path.name}")
            
        except Exception as e:
            print(f"Error processing {file_path.name}: {str(e)}")

if __name__ == "__main__":
    docs_dir = Path("data/Videos/Videos_summaries")
    index_documents(docs_dir) 